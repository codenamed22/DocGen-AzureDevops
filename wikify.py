import uuid
from azure.devops.connection import Connection
from msrest.authentication import BasicAuthentication
import azure.devops.released.wiki
import os
import tkinter
from tkinter.messagebox import showinfo, showerror
import tkinter.tix
from tkinter.constants import *
from tkinter import *
from tkinter import ttk
from tkinter import filedialog
import re
import pypandoc
import docx
import mistune
import custom_parser
from docx import Document
from docxcompose.composer import Composer
from docx import Document as Document_compose
import requests
import re
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import base64
import json
from docx.shared import Inches
from docx2python import docx2python

selected_items = []
wiki_content = []
wiki_client = ''
wiki_identifier = ''
project = ''


def store_cred(url, project, token):
    f = open("stored_cred.txt", "w+")
    f.write(url + "\n" + project + '\n' + token)
    f.close()


def print_wiki_page(_path, _wiki_content, project_name):
    global wiki_client
    global project
    wikipage = wiki_client.get_page(project=project_name, path=_path, wiki_identifier='SS365 ADO Wiki Templates',
                                    recursion_level='full', version_descriptor=None)
    modified_string = wikipage.page.path[1:]
    modified_string = modified_string.replace("/", ".")
    _wiki_content.append(modified_string)
    for subpage in wikipage.page.sub_pages:
        _wiki_content = print_wiki_page(subpage.path, _wiki_content, project_name)
    return _wiki_content


def fetch_wiki(mainroot, root, ent1, ent2, ent3, n):
    store_cred(ent1, ent2, ent3)
    # Fill in with your personal access token and org URL
    personal_access_token = ent3
    global organization_url
    organization_url = ent1

    # Create a connection to the org
    credentials = BasicAuthentication('', personal_access_token)
    connection = Connection(base_url=organization_url, creds=credentials)

    # Get a client (the "core" client provides access to projects, teams, etc)
    global wiki_client
    wiki_client = connection.clients.get_wiki_client()
    project_name = ent2
    global project
    project = ent2

    global wiki_content
    try:
        wiki_content = print_wiki_page('/', wiki_content, project_name)
    except:
        e = sys.exc_info()[0]
        showerror("Error", "There was an issue fetching the wiki\n %s" % e)
    wiki_content.sort(key=len)

    view = View(root)
    root.update()
    b1 = tkinter.Button(root, text="Generate",
                        command=lambda: generate(mainroot, root, view, n, organization_url, project_name,
                                                 personal_access_token), bg='white',
                        relief=tkinter.FLAT,
                        height=2, width=35)
    b1.grid(row=8, column=0, columnspan=2, pady=10)
    root.update()


class View(object):
    def __init__(self, root):
        self.root = root
        self.top = ''
        self.makeCheckList()

    def makeCheckList(self):
        self.top = tkinter.tix.Frame(self.root, relief=tkinter.FLAT, bd=1, bg='white')

        self.top.cl = tkinter.tix.CheckList(self.root, browsecmd=self.selectItem, width=400, height=300)
        self.top.grid(row=6, column=0, columnspan=2, sticky=EW)

        self.top.cl.grid(row=7, column=0, columnspan=2, sticky=EW)
        global wiki_content
        for page in wiki_content:
            if page != '':
                try:
                    self.top.cl.hlist.add(page, text=page.split('.')[-1])
                    self.top.cl.setstatus(page, "off")
                    self.top.cl.hlist.config(fg='grey', bg='white', pady='4', font='segoe')
                except:
                    print("Warning: Some issue with the checklist generation")
        self.top.cl.autosetmode()

    def selectItem(self, item):
        global selected_items
        if self.top.cl.getstatus(item) == 'on':
            if item not in selected_items:
                selected_items.append(item)
        else:
            if item in selected_items:
                selected_items.remove(item)


def generate(mainroot, root, view, n, organization_url, project_name, personal_access_token):
    global selected_items
    global wiki_client
    global project
    # try:
    b = ''
    for item in selected_items:
        new_item = '/' + item.replace(".", "/")
        wikipage = wiki_client.get_page(project=project, path=new_item, wiki_identifier='SS365 ADO Wiki Templates',
                                        recursion_level='full', version_descriptor=None, include_content=True)
        print(project + " " + new_item)
        b = b + wikipage.page.content

    f = open("imported1.md", "w")
    rb = b.replace("![image.png](/.", "image.png")
    rb = rb.replace("[[_TOC_]]", "")
    # rb = rb.replace("\t", "")
    # rb = rb.replace(r'\n', "")
    # rb = rb.replace("<", "")
    # rb = rb.replace(">", "")
    # rb = rb.replace("]", "")
    # rb = rb.replace("[", "")
    # rb = rb.replace("!", "")
    # rb = re.sub(r'https?:\/\/.*[\r\n]*', '', rb, flags=re.MULTILINE)
    f.write(rb)
    f.close()
    addWorkItem()
    directory = filedialog.asksaveasfilename(parent=root, defaultextension='.docx')
    pypandoc.convert_file("imported1.md", to='docx', format='markdown_mmd', outputfile="temp.docx")
    files_list = ['./doc_templates/' + n, 'temp.docx']
    mergeDoc(directory, files_list)
    addWorkItemLink(directory, organization_url, project_name, personal_access_token)
    insertImage(directory, organization_url, project_name, personal_access_token)
    showinfo("Success", "File has been generated!")
    # mainroot.destroy()
    # except:
    #     e = sys.exc_info()[0]
    #     showinfo("Error", "There was an issue generating the file\n %s" % e)
    #     root.destroy()


def upload_to_wiki(url, project, token, filename, wikiname):
    from azure.devops.connection import Connection
    from azure.devops.v5_1.wiki import WikiPageCreateOrUpdateParameters
    from msrest.authentication import BasicAuthentication
    personal_access_token = token
    organization_url = url
    # Create a connection to the org
    credentials = BasicAuthentication('', personal_access_token)
    connection = Connection(base_url=organization_url, creds=credentials)
    # convert the document into markk down.
    pypandoc.convert_file(filename, to='markdown_github', outputfile="foo.markdown")
    fddfilecheck = os.path.basename(filename)
    with open('foo.markdown', 'r') as fileForImages:
        Lines = fileForImages.readlines()
    updatedDocumentLines = list()
    # extract docx content
    doc_result = docx2python(filename)
    # extract images from document and add them to the git using the Rest API
    for key, val in doc_result.images.items():
        USER_PASS = "" + ":" + personal_access_token
        B64USERPASS = base64.b64encode(USER_PASS.encode()).decode()
        headers = {
            'Authorization': 'Basic' + B64USERPASS,
            'Content-Type': 'application/octet-stream'
        }
        imagedata = base64.b64encode(val)

        wikiUrl = url + "/" + project + "/_apis/wiki/wikis?api-version=5.0"
        wikiReq = requests.get(wikiUrl, headers=headers)
        wikiId = json.loads(wikiReq.content)["value"][0]["id"]

        URL = url + "/" + project + "/_apis/wiki/wikis/" + wikiId + "/attachments?name=" + str(
            uuid.uuid4()) + key + "&api-version=6.0"
        x = requests.put(URL, headers=headers, data=imagedata)
        statusCode = x.status_code
        if statusCode == 201:
            x = x.content
            values = json.loads(x)
            attachmentName = values["path"]
            for index, line in enumerate(Lines):
                imageMatch = re.search("<img src=\"media/" + key + "\"", line)
                if imageMatch is not None:
                    Lines.remove(line)
                    Lines.insert(index, str("![](" + attachmentName + ")"))

    with open('foo.markdown', 'w') as updatedFile:
        updatedFile.writelines(Lines)
        updatedFile.close()
    data = ''
    with open('foo.markdown', 'r') as file:
        data = file.read()
    file.close()
    client = connection.clients.get_wiki_client()
    project_name = project
    wikiPageCreateOrUpdateParameters = WikiPageCreateOrUpdateParameters()
    wikiPageCreateOrUpdateParameters.content = data
    a = client.create_or_update_page(project=project, path='/' + wikiname,
                                     wiki_identifier='SS365 ADO Wiki Templates',
                                     version=None, parameters=wikiPageCreateOrUpdateParameters)
    showinfo("Success", "File successfully pushed to wiki!")


# Function for opening the
# file explorer window
def browseFiles():
    global filename_string_value
    filename_string = filedialog.askopenfilename(initialdir="/",
                                                 title="Select the File",
                                                 filetypes=(("Word", "*.docx*"), ("Word", "*.doc*"))
                                                 )

    filename_string_value = filename_string


def main():
    import os
    file_list = os.listdir(os.path.dirname(os.path.abspath(__file__)) + "\doc_templates")
    root = tkinter.tix.Tk()

    root.wm_iconbitmap('logo.ico')
    root.title('Wikify')
    root.geometry("850x600")

    noteStyler = tkinter.ttk.Style()

    noteStyler.element_create('Plain.Notebook', "from", 'default')
    noteStyler.configure("TNotebook.Tab", padding=[10, 135], relief=tkinter.FLAT, background="darkred")
    noteStyler.map("TNotebook.Tab", background=[("selected", "darkred")], foreground=[("selected", "blue")]);
    noteStyler.configure("TNotebook", tabposition='wn')

    tabControl = tkinter.ttk.Notebook(root)
    tab1 = tkinter.ttk.Frame(tabControl, width=200, height=200)
    tab2 = tkinter.ttk.Frame(tabControl, width=200, height=200)
    tabControl.add(tab1, text='Wiki to Document')
    tabControl.add(tab2, text='Document to Wiki')

    tabControl.pack(expand=1, fill="both")
    lab = tkinter.Label(tab1, text="ADO URL")
    lab.grid(row=1, column=0, sticky=tkinter.W)

    ent1 = tkinter.Entry(tab1, width=100, relief=tkinter.FLAT)
    ent1.grid(row=1, column=1, pady=10, padx=5)
    lab2 = tkinter.Label(tab1, text="Project Name", anchor=tkinter.W)
    lab2.grid(row=2, column=0, sticky=tkinter.W)
    ent2 = tkinter.Entry(tab1, width=100, relief=tkinter.FLAT)
    ent2.grid(row=2, column=1, pady=10, padx=5)
    lab3 = tkinter.Label(tab1, text="Access token", anchor=tkinter.W)
    lab3.grid(row=3, column=0, sticky=tkinter.W)
    ent3 = tkinter.Entry(tab1, width=100, relief=tkinter.FLAT)
    ent3.grid(row=3, column=1, pady=10, padx=5)
    lab4 = tkinter.Label(tab1, text="Template", anchor=tkinter.W)
    lab4.grid(row=4, column=0, sticky=tkinter.W)
    n = StringVar(tab1)
    OPTIONS = file_list
    variable = StringVar(tab1)
    variable.set(OPTIONS[0])
    tkinter.ttk.Style().configure('TCombobox', padding=4, relief='flat', borderwidth=0, shiftrelief='flat')
    ent4 = tkinter.ttk.Combobox(tab1, width=100, textvariable=n)
    ent4['values'] = tuple(file_list)

    ent4.grid(row=4, column=1, pady=10, padx=5)
    f = open("stored_cred.txt", "r")
    ent1.insert(END, f.readline().rstrip())
    ent2.insert(END, f.readline().rstrip())
    ent3.insert(END, f.readline().rstrip())
    f.close()
    but = tkinter.Button(tab1, text="Fetch wiki",
                         command=lambda: fetch_wiki(root, tab1, ent1.get(), ent2.get(), ent3.get(), n.get()),
                         bg='white', relief=tkinter.FLAT,
                         height=2, width=35).grid(row=5, column=0, columnspan=2, pady=10)

    lab = tkinter.Label(tab2, text="ADO URL")
    lab.grid(row=1, column=0, sticky=tkinter.W)
    entry1 = tkinter.Entry(tab2, width=100, relief=tkinter.FLAT)
    entry1.grid(row=1, column=1, pady=10, padx=5)
    label2 = tkinter.Label(tab2, text="Project Name", anchor=tkinter.W)
    label2.grid(row=2, column=0, sticky=tkinter.W)
    entry2 = tkinter.Entry(tab2, width=100, relief=tkinter.FLAT)
    entry2.grid(row=2, column=1, pady=10, padx=5)
    label3 = tkinter.Label(tab2, text="Access token", anchor=tkinter.W)
    label3.grid(row=3, column=0, sticky=tkinter.W)
    entry3 = tkinter.Entry(tab2, width=100, relief=tkinter.FLAT)
    entry3.grid(row=3, column=1, pady=10, padx=5)
    label4 = tkinter.Label(tab2, text="File", anchor=tkinter.W)
    label4.grid(row=4, column=0, sticky=tkinter.W)
    # entry4 = tkinter.Entry(tab2, width=50, relief=tkinter.FLAT)
    # entry4.grid(row=4, column=1, pady=10, padx=5)
    # entry4.insert(4,filename_string_value)
    label5 = tkinter.Label(tab2, text="Wiki page name", anchor=tkinter.W)
    label5.grid(row=5, column=0, sticky=tkinter.W)
    entry5 = tkinter.Entry(tab2, width=100, relief=tkinter.FLAT)
    entry5.grid(row=5, column=1, pady=10, padx=5)
    f = open("stored_cred.txt", "r")
    entry1.insert(0, f.readline().rstrip())
    entry2.insert(0, f.readline().rstrip())
    entry3.insert(0, f.readline().rstrip())
    f.close()

    button = tkinter.Button(tab2, text="Push to wiki",
                            command=lambda: upload_to_wiki(entry1.get(), entry2.get(), entry3.get(),
                                                           filename_string_value,
                                                           entry5.get()),
                            bg='white', relief=tkinter.FLAT,
                            height=2, width=35)

    button.grid(row=6, column=0, columnspan=2, pady=10)
    # Create the root window

    button_explore = tkinter.Button(tab2,
                                    text="Select File",
                                    command=browseFiles, bg='white', relief=tkinter.FLAT,
                                    height=2, width=15)
    button_explore.grid(row=4, column=0, columnspan=2, pady=10)

    root.mainloop()


def insertImage(docpath, organization_url, project_name, personal_access_token):
    global wiki_client
    fdoc = docx.Document(docpath)
    USER_PASS = "" + ":" + personal_access_token
    B64USERPASS = base64.b64encode(USER_PASS.encode()).decode()
    headers = {'Authorization': 'Basic' + B64USERPASS}

    for p in fdoc.paragraphs:
        txt = p.text
        if "image.png" in txt:
            p.text = ""
            index1 = txt.find('attachments/')
            imageid = txt[index1 + 12:len(txt) - 1]

            wikiUrl = organization_url + "/" + project_name + "/_apis/wiki/wikis?api-version=5.0"
            wikiReq = requests.get(wikiUrl, headers=headers)
            wikiId = json.loads(wikiReq.content)["value"][0]["id"]

            URL = organization_url + "/" + project_name + "/_apis/git/repositories/" + wikiId + "/items//.attachments/" + imageid + "?versionType=Branch&versionOptions=None"
            x = requests.get(URL, headers=headers)
            b64image = base64.b64encode(x.content)
            from base64 import decodebytes
            with open("foo.png", "wb") as f:
                f.write(decodebytes(b64image))
            f.close()

            p.add_run().add_picture('foo.png', width=Inches(7))
    fdoc.save(docpath)


def addWorkItem():
    f = open("imported1.md", "r")
    Lines = f.readlines()
    a = ""
    for line in Lines:
        x = re.search("#{1}[0-9]", line)
        if (x is not None):
            a = a + "\n" + "workItem" + "#" + line[1:] + "\n"
        else:
            a += line
    f.close()

    fw = open("imported1.md", "w")
    fw.write(a)
    fw.close()


def addWorkItemLink(dir, organization_url, project_name, personal_access_token):
    global wiki_client
    fdoc = docx.Document(dir)

    for p in fdoc.paragraphs:
        txt = p.text
        if "workItem" in txt:
            wi = txt.find("#")
            num = txt[wi + 1:]
            p.text = ""

            USER_PASS = "" + ":" + personal_access_token
            B64USERPASS = base64.b64encode(USER_PASS.encode()).decode()
            headers = {'Authorization': 'Basic' + B64USERPASS}

            URL = organization_url + "/" + project_name + "/_apis/wit/workitems/" + num + "?api-version=6.0"
            x = requests.get(URL, headers=headers);
            statusCode = x.status_code
            if statusCode == 200:
                x = x.content
                values = json.loads(x)
                fields = values["fields"]
                title = fields["System.Title"]
                state = fields["System.State"]

                add_hyperlink(p, "Work Item " + num + " : " + title + " | " + state,
                              organization_url + "/_workitems/edit/" + num)
    fdoc.save(dir)


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def mergeDoc(filename_master, files_list):
    number_of_sections = len(files_list)
    master = Document_compose()
    composer = Composer(master)
    for i in range(0, number_of_sections):
        doc_temp = Document_compose(files_list[i])
        composer.append(doc_temp)
    composer.save(filename_master)


if __name__ == '__main__':
    main()
